from __future__ import annotations

import json
import os
import shutil
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
os.environ.setdefault("MPLBACKEND", "Agg")
os.environ.setdefault("MPLCONFIGDIR", str(ROOT / ".matplotlib"))

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


REPORT_DIR = ROOT / "reports" / "submission"
EXAMPLE_DIR = ROOT / "reports" / "E012_visual_examples"


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def fmt(value: float | int | str) -> str:
    try:
        numeric = float(value)
        if numeric.is_integer() and abs(numeric) < 1000:
            return str(int(numeric))
        return f"{numeric:.6f}"
    except (TypeError, ValueError):
        return str(value)


def add_table(document: Document, frame: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, col in enumerate(frame.columns):
        hdr[idx].text = str(col)
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = fmt(value)


def markdown_table(frame: pd.DataFrame) -> str:
    headers = list(frame.columns)
    rows = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for _, row in frame.iterrows():
        rows.append("| " + " | ".join(fmt(value) for value in row) + " |")
    return "\n".join(rows)


def make_score_chart(ranking: pd.DataFrame) -> Path:
    chart_path = REPORT_DIR / "final_score_chart.png"
    top = ranking.head(8).copy()
    plt.figure(figsize=(10, 5))
    plt.barh(top["exp_id"] + " " + top["model"], top["final_score"], color="#2f6f9f")
    plt.gca().invert_yaxis()
    plt.xlabel("Final Score")
    plt.title("Part 1 Model Ranking")
    plt.tight_layout()
    plt.savefig(chart_path, dpi=160)
    plt.close()
    return chart_path


def copy_score_tables() -> None:
    files = [
        ROOT / "runs" / "model_scoreboard.csv",
        ROOT / "reports" / "final_model_ranking.csv",
        ROOT / "reports" / "final_model_ranking.md",
        EXAMPLE_DIR / "tile_metrics_clean.csv",
        EXAMPLE_DIR / "tile_metrics_occluded.csv",
    ]
    for path in files:
        if path.exists():
            shutil.copy2(path, REPORT_DIR / path.name)


def build_content() -> tuple[str, pd.DataFrame, dict, dict]:
    ranking = pd.read_csv(ROOT / "reports" / "final_model_ranking.csv")
    scoreboard = pd.read_csv(ROOT / "runs" / "model_scoreboard.csv")
    clean = read_json(ROOT / "runs" / "E012_segformer" / "metrics_test_clean.json")
    occluded = read_json(ROOT / "runs" / "E012_segformer" / "metrics_test_occluded.json")
    winner = ranking.iloc[0]

    compact = ranking[
        [
            "rank",
            "exp_id",
            "model",
            "input_bands",
            "iou",
            "dice",
            "recall",
            "occlusion_recall",
            "connectivity_ratio",
            "final_score",
        ]
    ]

    md = f"""# Occlusion-Robust Road Extraction Experiment Report

## Problem Statement

The objective of Part 1 was to extract roads from Sentinel-2 satellite imagery even when roads are partially hidden by shadows, vegetation, vehicles, clouds, or urban clutter. The predicted road mask must be useful for later graph construction, so we evaluated not only pixel accuracy but also recall, occlusion robustness, and connectivity.

## Dataset Preparation

We prepared three Sentinel-2 AOIs using OSM road vectors as reference labels:

- `bengaluru_core`: train split, 124 tiles
- `hyderabad_mixed`: validation split, 16 tiles
- `bengaluru_edge`: held-out test split, 304 tiles

The final merged dataset has 444 image-mask tiles, including 364 road-positive tiles. Each tile is 256 x 256 pixels with stride 128. RGB imagery was used for baseline experiments, while RGBN imagery was used for the final occlusion-aware SegFormer runs.

## Experimental Pipeline

1. Download Sentinel-2 L2A bands B02, B03, B04, and B08.
2. Download OSM road vectors for the same AOIs.
3. Reproject OSM roads to the Sentinel-2 CRS.
4. Buffer OSM road lines by class-dependent road widths.
5. Rasterize roads into binary masks.
6. Tile images and masks into 256 x 256 samples.
7. Train all model families using the same train/validation/test protocol.
8. Evaluate every model on clean and synthetic-occluded test inputs.
9. Rank models using a final score that rewards road continuity and occlusion recovery.

## Models Tested

The experiment covered CNN baselines and transformer/attention-heavy models:

- U-Net
- ResNet U-Net
- UNet++
- DeepLabV3+
- SegFormer
- Swin-Unet
- TransUNet
- Mask2Former
- DINO/ViT Head
- DeepLabV3+ with RGBN and synthetic occlusion training
- SegFormer with RGBN and synthetic occlusion training
- Refined SegFormer with RGBN and synthetic occlusion training

## Metrics

The final score was:

`0.20 * IoU + 0.20 * Dice + 0.15 * Recall + 0.20 * Occlusion Recall + 0.15 * Connectivity Ratio + 0.10 * Relaxed IoU`

This score is appropriate for the hackathon problem because road masks must feed a graph-theoretic pipeline. A model with high visual overlap but broken road continuity is less useful than a model that preserves connected road structures.

## Final Ranking

{markdown_table(compact)}

## Best Model

The best model is **{winner['model']}**, experiment **{winner['exp_id']}**, using **{winner['input_bands']}** input.

Key held-out clean test metrics:

- IoU: {fmt(clean['iou'])}
- Dice: {fmt(clean['dice'])}
- Recall: {fmt(clean['recall'])}
- Relaxed IoU: {fmt(clean['relaxed_iou'])}
- Connectivity Ratio: {fmt(clean['connectivity_ratio'])}

Key synthetic-occluded test metrics:

- Occlusion Recall: {fmt(occluded['occlusion_recall'])}
- Occluded Recall: {fmt(occluded['recall'])}
- Occluded Connectivity Ratio: {fmt(occluded['connectivity_ratio'])}

Final combined score: **{fmt(winner['final_score'])}**

## Why This Model Won

The refined SegFormer model won because it combined long-range context with explicit occlusion-aware training. Roads are elongated structures, and transformer-style patch attention helps connect spatial evidence across a tile. Synthetic occlusion training forced the model to predict roads even when visible image evidence was degraded. Compared with the CNN baselines, the final model achieved much higher recall, occlusion recall, and connectivity ratio.

## Fit to the Original Problem

The problem asks for occlusion-robust road extraction that can support graph-theoretic criticality analysis. The final model fits this requirement because it produces road masks that are more complete under occlusion and more connected spatially. This makes it more suitable for the next phase: skeletonization, graph healing, bottleneck identification, and resilience simulation.

## Limitations and Failure Cases

The experiment used OSM-derived labels, so label noise and road-width assumptions affect the scores. Sentinel-2 has 10 m resolution, which limits fine-grained lane-level road extraction. Some failure cases still show overprediction in dense urban areas and missed narrow roads. These limitations are expected and should be improved with higher-resolution Resourcesat/Cartosat imagery and more AOIs.

## Conclusion

Based on the completed experiments, **E012: SegFormer + RGBN + synthetic occlusion training** is the best Part 1 model. It is the recommended segmentation model for the next graph reconstruction phase because it gives the strongest balance of road recovery, occlusion robustness, and connectivity.
"""
    return md, compact, clean, occluded


def build_docx(markdown_text: str, ranking: pd.DataFrame, clean: dict, occluded: dict, chart_path: Path) -> Path:
    doc_path = REPORT_DIR / "Occlusion_Robust_Road_Extraction_Final_Report.docx"
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = document.add_heading("Occlusion-Robust Road Extraction Experiment Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Part 1: Satellite Image Road Segmentation for Graph-Theoretic Urban Mobility Analysis")

    document.add_heading("1. Problem Statement", level=1)
    document.add_paragraph(
        "The goal was to extract roads from Sentinel-2 satellite imagery even when roads are partially hidden by "
        "vegetation, shadows, clouds, vehicles, or urban clutter. The predicted road mask must preserve road "
        "continuity so it can be converted into a graph in the next phase."
    )

    document.add_heading("2. Dataset Preparation", level=1)
    document.add_paragraph(
        "Three AOIs were prepared using Sentinel-2 RGB/NIR bands and OSM road vectors rasterized into binary masks."
    )
    dataset_table = pd.DataFrame(
        [
            ["bengaluru_core", "train", 124],
            ["hyderabad_mixed", "validation", 16],
            ["bengaluru_edge", "test", 304],
        ],
        columns=["AOI", "Split", "Tiles"],
    )
    add_table(document, dataset_table)

    document.add_heading("3. Experimental Method", level=1)
    for step in [
        "Downloaded Sentinel-2 L2A bands B02, B03, B04, and B08.",
        "Downloaded OSM road vectors for each AOI.",
        "Rasterized OSM roads into road masks using class-dependent road widths.",
        "Created 256 x 256 image-mask tiles with stride 128.",
        "Trained CNN and transformer segmentation models under the same protocol.",
        "Evaluated models on clean and synthetic-occluded test data.",
        "Ranked models using pixel, occlusion, and connectivity metrics.",
    ]:
        document.add_paragraph(step, style="List Number")

    document.add_heading("4. Final Score Formula", level=1)
    document.add_paragraph(
        "Final Score = 0.20*IoU + 0.20*Dice + 0.15*Recall + 0.20*Occlusion Recall + "
        "0.15*Connectivity Ratio + 0.10*Relaxed IoU"
    )
    document.add_paragraph(
        "This scoring formula is suitable because the downstream task requires a connected routable road network, "
        "not only a visually good pixel mask."
    )

    document.add_heading("5. Final Model Ranking", level=1)
    add_table(document, ranking)
    document.add_picture(str(chart_path), width=Inches(6.4))

    document.add_heading("6. Best Model and Metrics", level=1)
    document.add_paragraph("Best model: E012 SegFormer + RGBN + synthetic occlusion training.")
    best_metrics = pd.DataFrame(
        [
            ["Clean IoU", clean["iou"]],
            ["Clean Dice", clean["dice"]],
            ["Clean Recall", clean["recall"]],
            ["Clean Connectivity Ratio", clean["connectivity_ratio"]],
            ["Occlusion Recall", occluded["occlusion_recall"]],
            ["Occluded Connectivity Ratio", occluded["connectivity_ratio"]],
        ],
        columns=["Metric", "Value"],
    )
    add_table(document, best_metrics)

    document.add_heading("7. Visual Examples", level=1)
    for label, filename in [
        ("Strong clean predictions", "clean_strong_grid.png"),
        ("Typical clean predictions", "clean_typical_grid.png"),
        ("Failure cases", "clean_failure_grid.png"),
        ("Strong occluded predictions", "occluded_strong_grid.png"),
        ("Typical occluded predictions", "occluded_typical_grid.png"),
    ]:
        path = EXAMPLE_DIR / filename
        if path.exists():
            document.add_paragraph(label)
            document.add_picture(str(path), width=Inches(6.4))

    document.add_heading("8. Why SegFormer Fits the Problem", level=1)
    document.add_paragraph(
        "SegFormer fits the problem because roads are long continuous structures and transformer-style context "
        "helps reason across disconnected visible segments. RGBN input adds spectral information, while synthetic "
        "occlusion training directly teaches the model to recover roads under tree cover, shadows, and cloud-like "
        "degradation."
    )

    document.add_heading("9. Limitations", level=1)
    for item in [
        "OSM labels may be incomplete or misaligned.",
        "Sentinel-2 resolution is 10 m, so narrow roads are difficult to detect.",
        "Validation AOI is small, so future work should add more cities and Resourcesat/Cartosat imagery.",
        "Some predictions over-expand road regions in dense urban blocks.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("10. Conclusion", level=1)
    document.add_paragraph(
        "The completed experiments show that E012, a refined SegFormer using RGBN input and synthetic occlusion "
        "training, is the best model for Part 1. It gives the strongest balance of recall, occlusion recovery, "
        "and connectivity, making it the most suitable input to the graph reconstruction and criticality analysis phases."
    )

    document.save(doc_path)
    return doc_path


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    copy_score_tables()
    markdown_text, ranking, clean, occluded = build_content()
    (REPORT_DIR / "Occlusion_Robust_Road_Extraction_Final_Report.md").write_text(markdown_text, encoding="utf-8")
    chart_path = make_score_chart(ranking)
    docx_path = build_docx(markdown_text, ranking, clean, occluded, chart_path)
    print(REPORT_DIR / "Occlusion_Robust_Road_Extraction_Final_Report.md")
    print(docx_path)


if __name__ == "__main__":
    main()
