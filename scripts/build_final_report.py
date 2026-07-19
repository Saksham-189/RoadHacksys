from __future__ import annotations

import csv
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "reports" / "final_submission"
DOCX_PATH = OUTPUT / "Route_Resilience_Complete_Technical_Report.docx"
SUMMARY_PATH = OUTPUT / "submission_summary.json"
RESULTS_PATH = OUTPUT / "consolidated_results.csv"

NAVY = "17324D"
BLUE = "1F5C7A"
GREEN = "2F6F4E"
ORANGE = "C95D24"
LIGHT = "EEF2F4"
MID = "D5DEE3"
INK = "172027"
MUTED = "56636C"


def read_json(path: str) -> dict:
    return json.loads((ROOT / path).read_text(encoding="utf-8"))


def read_csv(path: str) -> list[dict]:
    with (ROOT / path).open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_properties = table._tbl.tblPr
    width = table_properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        table_properties.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), "9360")
    indent = table_properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(round(value * 1440)))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width_dxa = str(round(widths[index] * 1440))
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), width_dxa)


def set_run(run, size=10.5, bold=False, color=INK, italic=False) -> None:
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text: str, *, bold=False, italic=False, color=INK, after=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    set_run(paragraph.add_run(text), bold=bold, italic=italic, color=color)
    return paragraph


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.line_spacing = 1.10
        set_run(paragraph.add_run(item))


def add_steps(doc, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.line_spacing = 1.10
        set_run(paragraph.add_run(item))


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
            if index == 0
            else WD_ALIGN_PARAGRAPH.CENTER
        )
        set_run(paragraph.add_run(header), size=9, bold=True, color="FFFFFF")
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if row_index % 2:
                set_cell_shading(cells[index], "F7F9FA")
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if index == 0
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            set_run(paragraph.add_run(str(value)), size=8.6)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc, path: str, caption: str, width=6.35) -> None:
    image_path = ROOT / path
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(9)
    set_run(
        caption_paragraph.add_run(caption),
        size=8.5,
        italic=True,
        color=MUTED,
    )


def add_callout(doc, label: str, text: str, fill=LIGHT) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    set_run(paragraph.add_run(f"{label}: "), bold=True, color=NAVY)
    set_run(paragraph.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 15, 7),
        ("Heading 2", 13, BLUE, 11, 5),
        ("Heading 3", 11.5, GREEN, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)

    header = section.header.paragraphs[0]
    set_run(
        header.add_run("ROUTE RESILIENCE | TECHNICAL REPORT"),
        size=8,
        bold=True,
        color=MUTED,
    )
    footer = section.footer.paragraphs[0]
    set_run(footer.add_run("ISRO Hackathon Project  |  "), size=8, color=MUTED)
    add_page_number(footer)


def build_summary() -> dict:
    e013 = read_json("runs/e013_pretrained_segformer_rgbn/test_metrics.json")
    e013_training = read_json(
        "runs/e013_pretrained_segformer_rgbn/training_summary.json"
    )
    e012_ranking = read_csv("reports/final_model_ranking.csv")[0]
    part2 = read_json("runs/part25_consolidation/consolidation_summary.json")
    part3 = read_json("runs/part3/part3_summary.json")
    part4_baseline = read_json("runs/part4/baseline/baseline_summary.json")
    scenarios = read_csv("runs/part4/scenario_scoreboard.csv")
    dataset = read_csv("data/expanded_v2/final_manifest.csv")
    split_counts: dict[str, int] = {}
    aoi_counts: dict[str, int] = {}
    for row in dataset:
        split_counts[row["split"]] = split_counts.get(row["split"], 0) + 1
        aoi_counts[row["aoi"]] = aoi_counts.get(row["aoi"], 0) + 1
    worst = min(
        scenarios, key=lambda row: float(row["service_adjusted_resilience"])
    )
    return {
        "generated_on": date.today().isoformat(),
        "dataset": {
            "sensor": "Sentinel-2",
            "spatial_resolution_m": 10,
            "tiles": len(dataset),
            "aois": len(aoi_counts),
            "split_counts": split_counts,
        },
        "part1": {
            "current_model": "E013 pretrained SegFormer-B0",
            "input": "RGBN",
            "synthetic_occlusion_training": True,
            "test_tiles": e013["test_tiles"],
            "clean_iou": e013["clean"]["iou"],
            "clean_dice": e013["clean"]["dice"],
            "clean_recall": e013["clean"]["recall"],
            "occlusion_recall": e013["occluded"]["occlusion_recall"],
            "current_final_score": e013["final_score"],
            "parameters": e013_training["parameters"],
            "training_minutes": e013_training["elapsed_minutes"],
            "historical_v1_model": e012_ranking["exp_id"],
            "historical_v1_score": float(e012_ranking["final_score"]),
            "protocol_note": (
                "E012 and E013 use different datasets/evaluation protocols and "
                "their final scores are not directly comparable."
            ),
        },
        "part2": {
            "method": "B007 directional confidence-gated healing",
            "largest_component_after": part2["selected"][
                "largest_component_after"
            ],
            "route_success_rate": part2["selected"]["route_success_rate"],
            "false_bridge_rate": part2["selected"]["false_bridge_rate"],
            "node_coverage": part2["selected"]["node_coverage"],
            "coverage_gate_met": part2["coverage_gate_met"],
        },
        "part3": {
            "routing_nodes": part3["graph_preparation"]["routing_nodes"],
            "routing_edges": part3["graph_preparation"]["routing_edges"],
            "od_pairs": part3["demand"]["od_pairs"],
            "baseline_served_ratio": part3["gravity_baseline"][
                "served_demand_ratio"
            ],
            "top_flow_critical_node": part3["top_flow_critical_node"]["node_id"],
            "node_failure_disconnected_ratio": part3["top_flow_critical_node"][
                "disconnected_demand_ratio"
            ],
            "top_flow_critical_edge": (
                f"{part3['top_flow_critical_edge']['source']}-"
                f"{part3['top_flow_critical_edge']['target']}"
            ),
            "edge_failure_time_increase": part3["top_flow_critical_edge"][
                "travel_time_increase"
            ],
        },
        "part4": {
            "baseline_average_path_m": part4_baseline[
                "average_shortest_path_length_m"
            ],
            "scenarios": len(scenarios),
            "worst_scenario": worst["scenario_id"],
            "worst_disconnected_ratio": 1
            - float(worst["served_demand_ratio"]),
            "worst_resilience": float(worst["service_adjusted_resilience"]),
            "all_converged": all(
                float(row["msa_relative_gap"]) <= 0.001 for row in scenarios
            ),
        },
    }


def write_consolidated_results(summary: dict) -> None:
    rows = [
        {
            "part": "1",
            "stage": "Road segmentation",
            "selected_method": summary["part1"]["current_model"],
            "headline_metric": "Expanded-test final score",
            "value": summary["part1"]["current_final_score"],
            "status": "Current candidate",
        },
        {
            "part": "2",
            "stage": "Graph extraction and healing",
            "selected_method": summary["part2"]["method"],
            "headline_metric": "Route success rate",
            "value": summary["part2"]["route_success_rate"],
            "status": "Safe; coverage gate unmet",
        },
        {
            "part": "3",
            "stage": "Criticality and flow analysis",
            "selected_method": "Gravity demand + MSA flow-aware ablation",
            "headline_metric": "Critical-node disconnected demand",
            "value": summary["part3"]["node_failure_disconnected_ratio"],
            "status": "Complete",
        },
        {
            "part": "4",
            "stage": "Disruption simulation",
            "selected_method": "JSON scenario engine + exact MSA rerouting",
            "headline_metric": "Worst service-adjusted resilience",
            "value": summary["part4"]["worst_resilience"],
            "status": "Complete",
        },
    ]
    with RESULTS_PATH.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)


def build_report(summary: dict) -> None:
    doc = Document()
    configure_document(doc)

    doc.add_paragraph().paragraph_format.space_after = Pt(50)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(kicker.add_run("ISRO HACKATHON TECHNICAL SUBMISSION"), size=10, bold=True, color=ORANGE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(8)
    set_run(title.add_run("Route Resilience"), size=30, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(
        subtitle.add_run(
            "Occlusion-Robust Road Extraction and Graph-Theoretic "
            "Criticality Analysis for Urban Mobility"
        ),
        size=15,
        color=BLUE,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(38)
    add_callout(
        doc,
        "System objective",
        "Convert satellite imagery into a resilient urban transport graph, "
        "identify infrastructure whose failure matters most, and quantify "
        "the resulting mobility disruption.",
        "E8F0F3",
    )
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(
        metadata.add_run(f"Complete implementation report | {date.today():%d %B %Y}"),
        size=10,
        color=MUTED,
    )
    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    add_text(
        doc,
        "This project implements the complete four-part analytical chain required "
        "by the challenge: road segmentation under occlusion, graph reconstruction "
        "and healing, flow-aware criticality analysis, and predictive disruption "
        "simulation. The implementation operates on Sentinel-2 imagery and "
        "OpenStreetMap-derived supervision, then treats the extracted network as a "
        "transport graph for stress testing."
    )
    add_table(
        doc,
        ["Part", "Question answered", "Selected implementation", "Evidence"],
        [
            [
                "1",
                "Where are the roads despite occlusion?",
                "Pretrained SegFormer-B0, RGBN, synthetic occlusions",
                f"Score {summary['part1']['current_final_score']:.3f} on 304 expanded-test tiles",
            ],
            [
                "2",
                "How do fragmented pixels become a usable graph?",
                "Skeletonization, graph tracing, confidence-gated B007 healing",
                f"{100*summary['part2']['route_success_rate']:.2f}% route success",
            ],
            [
                "3",
                "Which nodes and links are critical?",
                "Degree baseline, betweenness reference, gravity demand, MSA ablation",
                f"{100*summary['part3']['node_failure_disconnected_ratio']:.2f}% demand lost at node 3900",
            ],
            [
                "4",
                "What happens when infrastructure fails?",
                "Composable scenario engine with preview and exact rerouting",
                f"Worst resilience {summary['part4']['worst_resilience']:.3f}",
            ],
        ],
        [0.45, 1.65, 2.7, 1.7],
    )
    add_callout(
        doc,
        "Principal result",
        "The compound flood scenario D002 is the most damaging preset disruption: "
        f"{100*summary['part4']['worst_disconnected_ratio']:.2f}% of estimated "
        f"demand is disconnected and service-adjusted resilience falls to "
        f"{summary['part4']['worst_resilience']:.3f}.",
        "FFF3EA",
    )

    doc.add_heading("1. Problem Definition", level=1)
    add_text(
        doc,
        "Urban mobility networks are vulnerable to clouds, vegetation, shadows, "
        "incomplete mapping, floods, accidents, construction and bridge closures. "
        "The challenge is not merely to segment visible roads. It is to recover a "
        "connected network, identify gatekeeper infrastructure, and estimate how "
        "the city responds when that infrastructure becomes unavailable."
    )
    doc.add_heading("1.1 Required outputs", level=2)
    add_bullets(
        doc,
        [
            "An occlusion-robust road mask inferred from satellite imagery.",
            "A topologically useful road graph with controlled false connections.",
            "Explainable critical-node and critical-edge rankings.",
            "Scenario-level resilience, disconnection, detour and rerouting maps.",
        ],
    )
    doc.add_heading("1.2 End-to-end system contract", level=2)
    add_steps(
        doc,
        [
            "Ingest Sentinel-2 RGB and near-infrared bands and align road labels.",
            "Infer a road probability surface and threshold it into a binary mask.",
            "Skeletonize and trace the mask into nodes, edges and geometries.",
            "Heal plausible gaps with confidence, distance and angle constraints.",
            "Assign relative capacity and synthetic gravity demand to the graph.",
            "Rank infrastructure by topology, flow and ablation damage.",
            "Apply disruption actions, reroute demand and calculate resilience.",
        ],
    )

    doc.add_heading("2. Data and Experimental Protocol", level=1)
    add_text(
        doc,
        f"The expanded manifest contains {summary['dataset']['tiles']:,} "
        f"Sentinel-2 tiles from {summary['dataset']['aois']} geographic AOIs. "
        "Every tile is 256 x 256 pixels at 10 m resolution. Geographic splits "
        "reduce tile leakage between training and final testing."
    )
    add_table(
        doc,
        ["Split", "Tiles", "Role"],
        [
            ["Train", f"{summary['dataset']['split_counts']['train']:,}", "Model fitting"],
            ["Validation", f"{summary['dataset']['split_counts']['val']:,}", "Early stopping and threshold selection"],
            ["Test", f"{summary['dataset']['split_counts']['test']:,}", "Held-out Bengaluru edge AOI"],
        ],
        [1.35, 1.0, 4.15],
    )
    add_callout(
        doc,
        "Resourcesat status",
        "The codebase retains a sensor-adapter path, but the measured experiments "
        "reported here use Sentinel-2 because no verified Resourcesat LISS-IV "
        "scene was available in the local dataset.",
    )

    doc.add_heading("3. Part 1: Occlusion-Robust Road Segmentation", level=1)
    add_text(
        doc,
        "The model study compared CNN baselines, transformer variants and "
        "synthetic-occlusion training. E013 is the current implementation: a "
        "pretrained SegFormer-B0 fine-tuned on RGBN inputs with tree, cloud, shadow, "
        "cutout, vehicle and haze-like input degradation."
    )
    add_table(
        doc,
        ["Metric", "Clean", "Occluded"],
        [
            ["IoU", f"{summary['part1']['clean_iou']:.3f}", "0.296"],
            ["Dice", f"{summary['part1']['clean_dice']:.3f}", "0.457"],
            ["Recall", f"{summary['part1']['clean_recall']:.3f}", "0.783"],
            ["Occlusion recall", f"{summary['part1']['clean_recall']:.3f}", f"{summary['part1']['occlusion_recall']:.3f}"],
            ["Final score", f"{summary['part1']['current_final_score']:.3f}", "Combined clean/occlusion/topology score"],
        ],
        [2.4, 1.4, 2.7],
    )
    add_callout(
        doc,
        "Protocol distinction",
        f"E012 achieved {summary['part1']['historical_v1_score']:.3f} in the earlier "
        "v1 experiment matrix. E013 achieved "
        f"{summary['part1']['current_final_score']:.3f} on the expanded 304-tile "
        "test protocol. These scores must not be interpreted as a direct regression "
        "because the data and evaluation pipelines differ.",
        "FFF3EA",
    )
    add_figure(
        doc,
        "runs/e013_pretrained_segformer_rgbn/examples_clean/bengaluru_edge_00303.png",
        "Figure 1. E013 clean-tile diagnostic: RGB input, OSM target, probability surface and thresholded overlay.",
    )
    add_text(
        doc,
        "The low operating threshold favours road recall and occlusion continuity, "
        "which is useful for downstream graph recovery but produces visible "
        "over-segmentation. This is the principal Part 1 failure mode and motivates "
        "confidence-gated graph healing rather than unconditional pixel bridging.",
        italic=True,
        color=MUTED,
    )

    doc.add_heading("4. Part 2: Road Graph Construction and Healing", level=1)
    add_text(
        doc,
        "A multi-stage hybrid is used because segmentation, topology and routing "
        "have different error structures. Morphological cleanup removes isolated "
        "noise; skeletonization reduces the road mask to centre lines; pixel-chain "
        "tracing creates graph nodes and edges; and B007 only bridges endpoints that "
        "satisfy distance, angular consistency and path-confidence gates."
    )
    add_table(
        doc,
        ["Measure", "Selected B007 result", "Interpretation"],
        [
            ["Largest component after healing", f"{100*summary['part2']['largest_component_after']:.2f}%", "Tile-level connectivity"],
            ["Route success rate", f"{100*summary['part2']['route_success_rate']:.2f}%", "Controlled-gap benchmark"],
            ["False-bridge rate", f"{100*summary['part2']['false_bridge_rate']:.2f}%", "Below the 10% safety gate"],
            ["Routing node coverage", f"{100*summary['part2']['node_coverage']:.2f}%", "Below the 70% target"],
        ],
        [2.25, 1.55, 2.7],
    )
    add_figure(
        doc,
        "runs/part2_healing/part2_comparison.png",
        "Figure 2. Comparative Part 2 healing experiments on controlled road gaps.",
    )
    add_callout(
        doc,
        "Safety decision",
        "No configuration reached 70% node coverage while remaining below 10% false "
        "bridges. The system therefore preserves the safest high-coverage graph and "
        "restricts routing analysis to its largest connected component.",
        "FFF3EA",
    )

    doc.add_heading("5. Part 3: Criticality and Flow-Aware Stress Analysis", level=1)
    add_text(
        doc,
        f"The routing graph contains {summary['part3']['routing_nodes']:,} nodes and "
        f"{summary['part3']['routing_edges']:,} edges. Relative road width defines "
        "three capacity classes. A reproducible gravity model generates "
        f"{summary['part3']['od_pairs']:,} OD pairs, and MSA assigns demand under "
        "BPR congestion costs."
    )
    add_table(
        doc,
        ["Method", "Role", "Result"],
        [
            ["Degree centrality", "Simple explainable baseline", "Ranks locally connected junctions"],
            ["Betweenness", "Mandatory gatekeeper reference", "Finds shortest-path intermediaries"],
            ["Flow-aware ablation", "Advanced final ranking", "Measures rerouting and disconnected demand"],
            ["Node 3900 removal", "Top critical node", f"{100*summary['part3']['node_failure_disconnected_ratio']:.2f}% demand disconnected"],
            [f"Edge {summary['part3']['top_flow_critical_edge']}", "Top critical edge", f"{100*summary['part3']['edge_failure_time_increase']:.2f}% mean-time increase"],
        ],
        [1.65, 2.1, 2.75],
    )
    add_figure(
        doc,
        "runs/part3/node_criticality_map.png",
        "Figure 3. Flow-aware node criticality over the selected routing component.",
    )
    add_figure(
        doc,
        "runs/part3/resilience_curves.png",
        "Figure 4. Progressive network degradation under random and targeted node-removal strategies.",
    )

    doc.add_heading("6. Part 4: Disruption Simulation and Resilience", level=1)
    add_text(
        doc,
        "The reusable SimulationEngine accepts node, edge, capacity, circle and "
        "polygon actions. Preview mode reuses baseline edge costs for rapid feedback; "
        "exact mode reruns MSA to convergence and reconstructs every surviving OD "
        "route. The baseline graph is copied for each scenario and never mutated."
    )
    add_text(
        doc,
        "Official resilience = served demand ratio x demand-weighted baseline path "
        "length / demand-weighted disrupted path length. The canonical path ratio "
        "and travel-time change remain separately visible."
    )
    scenarios = read_csv("runs/part4/scenario_scoreboard.csv")
    scenario_rows = []
    for row in scenarios:
        scenario_rows.append(
            [
                row["scenario_id"],
                row["hazard_type"].replace("_", " "),
                f"{100*(1-float(row['served_demand_ratio'])):.2f}%",
                f"{100*float(row['affected_demand_ratio']):.2f}%",
                f"{float(row['service_adjusted_resilience']):.3f}",
            ]
        )
    add_table(
        doc,
        ["ID", "Hazard", "Disconnected", "Affected", "Resilience"],
        scenario_rows,
        [0.55, 2.2, 1.25, 1.15, 1.35],
    )
    add_figure(
        doc,
        "runs/part4/resilience_comparison.png",
        "Figure 5. Service-adjusted resilience across the nine preset disruption scenarios.",
    )
    add_figure(
        doc,
        "runs/part4/scenario_impact_map.png",
        "Figure 6. Five-hundred-metre affected-demand cells for the worst preset scenario, D002.",
        width=5.6,
    )

    doc.add_heading("7. Integrated Findings", level=1)
    add_bullets(
        doc,
        [
            "Occlusion robustness is primarily a continuity problem, not only a pixel-overlap problem.",
            "Aggressive graph bridging improves connectivity but can create unsafe false links; B007 is selected by a safety gate rather than by raw healing score.",
            "Degree alone is insufficient: the most damaging assets emerge when topology, relative capacity, demand and rerouting are evaluated together.",
            "Node failures mainly remove service, while some edge failures preserve connectivity but create extreme congestion and rerouting burden.",
            "Targeted critical-node removal is substantially more damaging than matched random failure, supporting the value of the criticality ranking.",
        ],
    )
    add_callout(
        doc,
        "Urban-planning interpretation",
        "The system converts a satellite-derived map into a question planners can "
        "act on: if this junction or corridor fails, what fraction of estimated "
        "mobility demand loses service, how much longer do surviving routes become, "
        "and where does rerouted burden appear?",
        "E8F0F3",
    )

    doc.add_heading("8. Limitations and Responsible Interpretation", level=1)
    add_bullets(
        doc,
        [
            "Sentinel-2 at 10 m cannot resolve lane counts; capacity classes are relative corridor-width proxies.",
            "Gravity demand is graph-derived and reproducible, but it is not observed traffic demand.",
            "OSM provides supervision and controlled validation; incomplete OSM labels can penalize valid inferred roads.",
            "The selected safe graph covers 63.87% of estimated nodes, so citywide coverage remains an unmet gate.",
            "E013 uses a recall-oriented threshold and visibly over-segments some dense urban textures.",
            "Resourcesat cross-sensor evaluation remains pending verified LISS-IV data.",
        ],
    )

    doc.add_heading("9. Reproducibility", level=1)
    add_text(doc, "Primary commands:")
    commands = [
        r".\.venv-win\Scripts\python.exe train.py --config configs\e013_pretrained_segformer_rgbn.yaml",
        r".\.venv-win\Scripts\python.exe generate_part2_graph.py",
        r".\.venv-win\Scripts\python.exe consolidate_part2_graph.py --config configs\part3_flow.yaml",
        r".\.venv-win\Scripts\python.exe run_part3.py --config configs\part3_flow.yaml",
        r".\.venv-win\Scripts\python.exe run_phase4_scenarios.py --config configs\phase4.yaml",
        r".\.venv-win\Scripts\python.exe -m pytest -q",
    ]
    for command in commands:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(4)
        set_run(paragraph.add_run(command), size=8.5, color=NAVY)
    add_text(
        doc,
        "The current repository passes 22 automated tests covering graph "
        "preparation, centrality, assignment, scenario validation, resilience "
        "bounds, reprojection, artifact schemas and end-to-end integration.",
    )

    doc.add_heading("10. Conclusion", level=1)
    add_text(
        doc,
        "The project demonstrates a complete decision pipeline rather than an "
        "isolated segmentation model. Satellite imagery supplies the road evidence; "
        "graph healing converts it into a routing domain; flow-aware ablation "
        "identifies gatekeeper infrastructure; and the disruption engine translates "
        "those rankings into demand loss, detour, congestion and geographic impact. "
        "The result fits the problem because it directly answers the planning "
        "question: which failures matter most, where do their consequences appear, "
        "and how resilient is the remaining urban network?"
    )

    doc.save(DOCX_PATH)


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    summary = build_summary()
    SUMMARY_PATH.write_text(json.dumps(summary, indent=2), encoding="utf-8")
    write_consolidated_results(summary)
    build_report(summary)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
